from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.lang import Builder
from docx import Document

def pricecalculate(price, width, length, shtuka):
    price = float(price)
    width = float(width)
    length = float(length)

    if shtuka == "плитка":
        result = price * width * length / 0.225

    elif shtuka == "ламинат":
        result = price * width * length  / 2

    elif shtuka == "обои":
        result = price * width * length  / 5

    return result

def kolcalculate(price, width, length, shtuka):
    price = float(price)
    width = float(width)
    length = float(length)
    if shtuka == "плитка":
        result = width * length / 0.225

    elif shtuka == "ламинат":
        result = width * length  / 2

    elif shtuka == "обои":
        result = width * length  / 5

    return result


class MainWidget(BoxLayout):
    def on_button_press(self):
        price = self.ids.price.text
        width = self.ids.width.text
        length = self.ids.length.text
        shtuka = self.ids.shtuka.text

        result1 = pricecalculate(price, width, length, shtuka)
        result2 = kolcalculate(price, width, length, shtuka)
        self.ids.result_label.text = f"Стоимость: {result1}, Требуемое кол-во: {result2}"
        self.savedoc(shtuka, result1, result2)
        
    def savedoc(self, shtuka, cost, quantity):
        doc = Document()
        doc.add_heading("Отчёт по расчёту материалов", level=1)
        doc.add_paragraph(f"Материал: {shtuka}")
        doc.add_paragraph(f"Требуемое количество: {quantity} шт.")
        doc.add_paragraph(f"Общая стоимость: {cost} руб.")
        doc.save("report.docx")

class MyApp(App):
    def build(self):
        Builder.load_file('main.kv')
        return MainWidget()

if __name__ == '__main__':
    MyApp().run()

