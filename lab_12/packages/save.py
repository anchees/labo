from docx import Document

def save_doc(material, result):
    doc = Document()
    doc.add_heading("Отчёт по расчёту материалов", level=1)
    doc.add_paragraph(f"Материал: {material}")
    doc.add_paragraph(f"Количество: {result['quantity']} шт.")
    doc.add_paragraph(f"Стоимость: {result['cost']} руб.")
    
    doc.save("report.docx")
